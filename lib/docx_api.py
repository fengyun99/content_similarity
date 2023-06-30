#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time     : 2023/6/30 14:24
# @Author   : FengYun
# @File     : docx_api.py
# @Software : PyCharm
import os
import re
import time
import sys

import docx
import pythoncom
import win32com
from bs4 import BeautifulSoup
from pydocx import PyDocX

from win32com.client import Dispatch
import collections

# python版本太高，需要添加abc，来找到以前的模块
if sys.version_info >= (3, 9):
    collections.Hashable = collections.abc.Hashable

# 默认style
init_style = """
* { margin: 0;padding: 0;}
body { padding: 10px !important;}
table td {border-style: solid;}
"""


class WordTransformApi:

    # 提取docx的纯文本 后保存到新的docx文件中
    def word_extract_text(self, filepath):
        if filepath is None:
            raise ValueError('文件不存在')

        # if not filepath.endswith('.docx'):
        #     raise ValueError('只支持docx文件')

        try:
            # linux 不支持该库
            # 调用doc转为docx在读取内容，win32
            if filepath.endswith('.doc'):
                html = self.doc_to_html(filepath)
            else:
                # 先将word转html
                html = PyDocX.to_html(filepath)

            # # 先将docx转html
            # html = PyDocX.to_html(filepath)
            # 将文本域中的空格替换为&nbsp;
            html = html.replace(" ", "&nbsp;")

            soup = BeautifulSoup(html, 'lxml')
            meta = soup.new_tag(
                'meta', content='width=device-width, initial-scale=1.0, maximum-scale=1.0, user-scalable=0;')
            soup.head.insert(1, meta)
            # 删除style
            soup.style.clear()
            soup.style.append(init_style)

            # 清理p标签里的span标签（解决后续因去标签时导致换行）
            for p in soup.find_all('p'):
                # 获取<p>标签中的文本内容
                text = p.getText()
                # 删除<span>标签及其属性
                for span in p.find_all('span'):
                    span.decompose()
                # 将<p>标签中的文本内容替换为去除了<span>标签及其属性的文本内容
                p.string = text
            # 格式化html，自动补全代码
            formatted_html = soup.prettify()

            # 转txt
            soup2 = BeautifulSoup(formatted_html, "html.parser")

            # 删除 script 和 style 标签
            for script in soup2(["style", "script", "head", "title", "meta", "[document]"]):
                script.extract()
            # 获取纯文本
            text = soup2.get_text()
            # 去掉多余的空格和空行
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text = '\n'.join(chunk for chunk in chunks if chunk)

            # 定义要删除的正则表达式(-1- 这样的页脚)
            pattern = re.compile(r'-\s\d+\s-')
            new_text = "\n".join(line for line in re.split(r"\n", text) if not pattern.search(line))
            new_text_list = new_text.split("\n")

            new_doc = docx.Document()
            # 添加文本到新的文档
            for new_text_tmp in new_text_list:
                if len(new_text_tmp) > 3 and self.is_contain_chinese(new_text_tmp):
                    new_text_tmp = new_text_tmp.replace(' ', '')
                    new_doc.add_paragraph(new_text_tmp)

            # 保存
            ext = filepath.split(".")[-1]
            new_path = filepath[:-len(ext) - 1] + "_transferred.docx"

            # 存在则删除
            if new_path and os.path.exists(new_path):
                os.remove(new_path)

            new_doc.save(new_path)
            print('文件转化成功: ' + filepath)
            return new_path
        except Exception as e:
            print('文档转化失败')
            print(e)
            raise Exception('文档转化失败')

    # linux 不支持该库
    # 将特殊格式的doc文件转为docx，然后提取html，然后删除临时文件
    def doc_to_html(self, filepath):
        new_path = None
        doc = None
        w = None
        try:
            w = win32com.client.Dispatch('Word.Application')
            w.Visible = 0
            w.DisplayAlerts = 0
            doc = w.Documents.Open(filepath)
            # 这里必须要绝对地址,保持和doc路径一致
            ext = filepath.split(".")[-1]
            new_path = filepath[:-len(ext) - 1] + "_transferred.docx"
            # 暂停1s，否则会出现-2147352567,错误
            time.sleep(1)
            if new_path and os.path.exists(new_path):
                os.remove(new_path)

            doc.SaveAs(new_path, 12, False, "", True, "", False, False, False, False)
            html = PyDocX.to_html(new_path)  # 出问题
            return html
        except Exception as e:
            print('文档doc转html失败')
            print(e)
            raise Exception('文档doc转html失败')
        finally:
            if doc is not None:
                doc.Close()

            if w is not None:
                w.Quit()

            # 删除临时文件
            if new_path and os.path.exists(new_path):
                os.remove(new_path)

    # 检查整个字符串是否包含中文
    def is_contain_chinese(self, text):
        for ch in text:
            if u'\u4e00' <= ch <= u'\u9fff':
                return True

        return False


class Docx:

    def read_content(self, docx_path):
        # 打开Word文档
        document = docx.Document(docx_path)
        paragraphs = document.paragraphs

        # 判断文件是否需要转换处理
        count = 0
        for paragraph in paragraphs:
            if len(paragraph.text) > 0:
                count += 1
        # 将原文档转换
        if count < 3:
            dtd = WordTransformApi()
            file_path = dtd.word_extract_text(docx_path)
            document = docx.Document(file_path)
            paragraphs = document.paragraphs
            if file_path.endswith("_transferred.docx") and os.path.exists(file_path):
                os.remove(file_path)
        message = ''
        for paragraph in paragraphs:
            message += paragraph.text.rstrip() + '\n'

        return message

    # 转换doc到docx
    def doc_to_docx(self, file_path):
        # 在需要使用 COM 的代码之前调用 CoInitialize 进行初始化
        pythoncom.CoInitialize()
        # 调用需要使用 COM 的函数或操作
        doc2docx = WordTransformApi()
        doc2docx.word_extract_text(file_path)
        # 在程序结束时，调用 CoUninitialize 释放 COM 资源
        pythoncom.CoUninitialize()
        # # 删除doc文件，重命名docx文件
        # os.remove(file_path)
        # print(f'删除原文件{file_path}')
        # test_transferred.docx
        old_name = file_path.split(".")[-2] + "_transferred" + ".docx"
        return old_name

    # 装载上传文件doc
    def loadDoc(self, file_name, file_path):
        message = self.read_content(file_path)
        name = "tmp/" + file_name.split("/")[-1] + ".txt"
        # print(name)
        # 确保文件夹存在
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        if not os.path.exists(file_path):
            # 文件不存在，创建文件
            open(f"{name}", "w", encoding="utf-8").close()
            with open(f"{name}", "w", encoding="utf-8") as f:
                # 进行下一步操作，比如写入内容等
                f.write(message + "\n")
        else:
            # 文件已存在，进行重命名操作
            index = 1
            while os.path.exists(name):
                # 生成新的文件名，添加序号
                name = "tmp/" + f"{file_name.split('/')[-1]}_{index}.txt"
                index += 1
                print(name)

            with open(name, "w", encoding="utf-8") as f:
                # 进行下一步操作，比如写入内容等
                f.write(message + "\n")

        return name